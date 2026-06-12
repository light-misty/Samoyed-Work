"""Word 文档处理器
基于 python-docx 实现 Word 文档的读取、转换、分析
精简版：仅支持 read/convert/analyze 操作
"""

import os
import html
import logging

from docx import Document


class WordHandler:
    """Word (.docx) 文档处理器（精简版：仅支持 read/convert/analyze）"""

    logger = logging.getLogger(__name__)

    # ------------------------------------------------------------------ #
    #  读取
    # ------------------------------------------------------------------ #

    def read(self, params: dict) -> dict:
        """读取 Word 文档内容"""
        path = params.get("path", "")
        if not path:
            self.logger.error("read: 缺少文件路径")
            return {"error": "缺少文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("read: 开始读取 Word 文档, path=%s", path)

        doc = Document(path)

        paragraphs = []
        for para in doc.paragraphs:
            para_info = {
                "text": para.text,
                "style": para.style.name if para.style else None,
            }
            paragraphs.append(para_info)

        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        props = {
            "title": doc.core_properties.title or "",
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
        }

        self.logger.info("read: Word 文档读取完成, path=%s, 段落数=%d, 表格数=%d", path, len(paragraphs), len(tables))
        return {
            "paragraphs": paragraphs,
            "tables": tables,
            "properties": props,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
        }

    # ------------------------------------------------------------------ #
    #  格式转换
    # ------------------------------------------------------------------ #

    def convert(self, params: dict) -> dict:
        """格式转换"""
        path = params.get("path", "")
        output_path = params.get("output_path", "")
        target_format = params.get("format", "md")
        if not path:
            self.logger.error("convert: 缺少源文件路径")
            return {"error": "缺少源文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("convert: 开始格式转换, path=%s, format=%s", path, target_format)

        doc = Document(path)

        if target_format in ("md", "markdown"):
            content = self._convert_to_markdown(doc)
        elif target_format == "txt":
            content = "\n".join(para.text for para in doc.paragraphs)
        elif target_format == "pdf":
            content = self._convert_to_pdf(doc, output_path or os.path.splitext(path)[0] + ".pdf")
            content = None
        else:
            self.logger.error("convert: 不支持的目标格式: %s", target_format)
            return {"error": f"不支持的目标格式: {target_format}"}

        if content is None:
            self.logger.info("convert: 格式转换完成, output_path=%s, format=%s", output_path, target_format)
            return {
                "path": output_path,
                "format": target_format,
                "message": f"已转换为 {target_format} 格式",
            }
        elif output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            self.logger.info("convert: 格式转换完成, output_path=%s, format=%s", output_path, target_format)
            return {
                "path": output_path,
                "format": target_format,
                "message": f"已转换为 {target_format} 格式",
            }
        else:
            return {
                "content": content,
                "format": target_format,
            }

    def _convert_to_markdown(self, doc: Document) -> str:
        """将 Word 文档内容转换为 Markdown"""
        lines = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text
            if not text.strip():
                continue
            if "Heading 1" in style:
                lines.append(f"# {text}")
            elif "Heading 2" in style:
                lines.append(f"## {text}")
            elif "Heading 3" in style:
                lines.append(f"### {text}")
            elif "Heading 4" in style:
                lines.append(f"#### {text}")
            elif "List" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)

        for table in doc.tables:
            lines.append("")
            for i, row in enumerate(table.rows):
                row_text = "| " + " | ".join(cell.text for cell in row.cells) + " |"
                lines.append(row_text)
                if i == 0:
                    lines.append("| " + " | ".join("---" for _ in row.cells) + " |")
            lines.append("")

        return "\n\n".join(lines)

    def _convert_to_pdf(self, doc: Document, output_path: str) -> None:
        """将 Word 文档内容转换为 PDF（使用 reportlab 渲染）"""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

        from handlers.font_utils import register_chinese_font
        font_name = register_chinese_font()

        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

        doc_pdf = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontName=font_name, fontSize=24, spaceAfter=30)
        heading_style = ParagraphStyle("CustomHeading", parent=styles["Heading2"], fontName=font_name, fontSize=16, spaceAfter=12)
        body_style = ParagraphStyle("CustomBody", parent=styles["Normal"], fontName=font_name, fontSize=12, leading=20, spaceAfter=10)

        elements = []
        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text
            if not text.strip():
                continue
            if "Heading 1" in style:
                elements.append(Paragraph(html.escape(text), heading_style))
            elif "Heading 2" in style:
                elements.append(Paragraph(html.escape(text), heading_style))
            elif "Heading 3" in style:
                elements.append(Paragraph(html.escape(text), body_style))
            elif "Title" in style:
                elements.append(Paragraph(html.escape(text), title_style))
            else:
                elements.append(Paragraph(html.escape(text), body_style))
            elements.append(Spacer(1, 0.3 * cm))

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, "#999999"),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                ]))
                elements.append(t)
                elements.append(Spacer(1, 0.5 * cm))

        doc_pdf.build(elements)

    # ------------------------------------------------------------------ #
    #  分析
    # ------------------------------------------------------------------ #

    def analyze(self, params: dict) -> dict:
        """分析 Word 文档"""
        path = params.get("path", "")
        if not path:
            self.logger.error("analyze: 缺少文件路径")
            return {"error": "缺少文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("analyze: 开始分析 Word 文档, path=%s", path)

        doc = Document(path)

        total_chars = sum(len(p.text) for p in doc.paragraphs)
        total_words = sum(len(p.text.split()) for p in doc.paragraphs)
        heading_count = sum(
            1 for p in doc.paragraphs
            if p.style and ("Heading" in p.style.name or "Title" in p.style.name)
        )

        headings = []
        for para in doc.paragraphs:
            if para.style and ("Heading" in para.style.name or "Title" in para.style.name):
                level = 1
                try:
                    level = int(para.style.name.replace("Heading ", ""))
                except ValueError:
                    if "Title" in para.style.name:
                        level = 0
                headings.append({"level": level, "text": para.text})

        self.logger.info("analyze: Word 文档分析完成, path=%s, 段落数=%d, 标题数=%d", path, len(doc.paragraphs), heading_count)
        return {
            "file_size": os.path.getsize(path),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "total_chars": total_chars,
            "total_words": total_words,
            "heading_count": heading_count,
            "headings": headings,
            "properties": {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            },
        }

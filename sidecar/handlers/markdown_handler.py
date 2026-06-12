"""Markdown 文档处理器
实现 Markdown 文档的读取、转换、分析
精简版：仅支持 read/convert/analyze 操作
"""

import os
import re
import html
import logging


class MarkdownHandler:
    """Markdown (.md) 文档处理器（精简版，仅支持 read/convert/analyze）"""

    logger = logging.getLogger(__name__)

    def read(self, params: dict) -> dict:
        """读取 Markdown 文档"""
        path = params.get("path", "")
        if not path:
            self.logger.error("read: 缺少文件路径")
            return {"error": "缺少文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("read: 开始读取 Markdown 文档, path=%s", path)

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        # 解析标题结构
        headings = []
        for match in re.finditer(r"^(#{1,6})\s+(.+)$", content, re.MULTILINE):
            level = len(match.group(1))
            text = match.group(2).strip()
            headings.append({"level": level, "text": text})

        self.logger.info("read: Markdown 文档读取完成, path=%s, 标题数=%d", path, len(headings))
        return {
            "content": content,
            "headings": headings,
            "heading_count": len(headings),
            "line_count": content.count("\n") + 1,
            "char_count": len(content),
        }

    def convert(self, params: dict) -> dict:
        """格式转换

        params:
            path: 源文件路径
            output_path: 输出文件路径
            format: 目标格式（html, txt）
        """
        path = params.get("path", "")
        output_path = params.get("output_path", "")
        target_format = params.get("format", "html")
        if not path:
            self.logger.error("convert: 缺少源文件路径")
            return {"error": "缺少源文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("convert: 开始格式转换, path=%s, format=%s", path, target_format)

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        if target_format in ("html", "htm"):
            # Markdown -> HTML: 简单转换，处理标题、段落、列表、代码块、粗体/斜体
            lines = content.split("\n")
            html_parts = []
            in_code_block = False
            in_list = False

            for line in lines:
                # 代码块处理
                if line.strip().startswith("```"):
                    if in_code_block:
                        html_parts.append("</code></pre>")
                        in_code_block = False
                    else:
                        lang = line.strip()[3:].strip()
                        html_parts.append(f'<pre><code class="language-{lang}">' if lang else "<pre><code>")
                        in_code_block = True
                    continue

                if in_code_block:
                    html_parts.append(html.escape(line))
                    continue

                # 标题处理
                heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
                if heading_match:
                    if in_list:
                        html_parts.append("</ul>")
                        in_list = False
                    level = len(heading_match.group(1))
                    text = heading_match.group(2).strip()
                    html_parts.append(f"<h{level}>{html.escape(text)}</h{level}>")
                    continue

                # 列表处理
                list_match = re.match(r"^[-*+]\s+(.+)$", line)
                if list_match:
                    if not in_list:
                        html_parts.append("<ul>")
                        in_list = True
                    text = list_match.group(1)
                    html_parts.append(f"<li>{html.escape(text)}</li>")
                    continue

                # 空行
                if not line.strip():
                    if in_list:
                        html_parts.append("</ul>")
                        in_list = False
                    continue

                # 普通段落
                if in_list:
                    html_parts.append("</ul>")
                    in_list = False
                # 处理行内格式
                text = html.escape(line)
                text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
                text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
                text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
                html_parts.append(f"<p>{text}</p>")

            if in_list:
                html_parts.append("</ul>")
            if in_code_block:
                html_parts.append("</code></pre>")

            result_content = "\n".join(html_parts)

        elif target_format == "txt":
            # Markdown -> 纯文本: 移除 Markdown 标记
            result_content = content
            # 移除标题标记
            result_content = re.sub(r"^#{1,6}\s+", "", result_content, flags=re.MULTILINE)
            # 移除粗体/斜体标记
            result_content = re.sub(r"\*\*(.+?)\*\*", r"\1", result_content)
            result_content = re.sub(r"\*(.+?)\*", r"\1", result_content)
            # 移除行内代码标记
            result_content = re.sub(r"`(.+?)`", r"\1", result_content)
            # 移除代码块标记
            result_content = re.sub(r"```\w*\n?", "", result_content)
            # 移除列表标记
            result_content = re.sub(r"^[-*+]\s+", "", result_content, flags=re.MULTILINE)
            # 移除链接，保留文本
            result_content = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", result_content)
            # 移除图片标记
            result_content = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", result_content)

        elif target_format in ("docx", "doc"):
            # Markdown -> Word: 使用 python-docx 生成 Word 文档
            from docx import Document
            from docx.shared import Pt

            if not output_path:
                output_path = os.path.splitext(path)[0] + ".docx"

            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

            doc = Document()
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i]
                stripped = line.strip()

                # 标题
                heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
                if heading_match:
                    level = len(heading_match.group(1))
                    text = heading_match.group(2).strip()
                    doc.add_heading(text, level=min(level, 9))
                    i += 1
                    continue

                # 代码块
                if stripped.startswith("```"):
                    code_lines = []
                    i += 1
                    while i < len(lines) and not lines[i].strip().startswith("```"):
                        code_lines.append(lines[i])
                        i += 1
                    i += 1  # 跳过结束的 ```
                    code_text = "\n".join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = doc.styles["Normal"]
                    for run in p.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(9)
                    continue

                # 列表
                list_match = re.match(r"^[-*+]\s+(.+)$", stripped)
                if list_match:
                    text = list_match.group(1)
                    doc.add_paragraph(text, style="List Bullet")
                    i += 1
                    continue

                # 有序列表
                olist_match = re.match(r"^\d+\.\s+(.+)$", stripped)
                if olist_match:
                    text = olist_match.group(1)
                    doc.add_paragraph(text, style="List Number")
                    i += 1
                    continue

                # 空行
                if not stripped:
                    i += 1
                    continue

                # 普通段落
                doc.add_paragraph(stripped)
                i += 1

            doc.save(output_path)
            result_content = None  # docx 已写入文件

        elif target_format == "pdf":
            # Markdown -> PDF: 使用 reportlab 渲染
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

            # 注册中文字体
            from handlers.font_utils import register_chinese_font
            font_name = register_chinese_font()

            if not output_path:
                output_path = os.path.splitext(path)[0] + ".pdf"

            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)

            doc_pdf = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle("CustomTitle", parent=styles["Title"], fontName=font_name, fontSize=24, spaceAfter=30)
            h1_style = ParagraphStyle("H1", parent=styles["Heading1"], fontName=font_name, fontSize=20, spaceAfter=15)
            h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontName=font_name, fontSize=16, spaceAfter=12)
            h3_style = ParagraphStyle("H3", parent=styles["Heading3"], fontName=font_name, fontSize=14, spaceAfter=10)
            body_style = ParagraphStyle("CustomBody", parent=styles["Normal"], fontName=font_name, fontSize=12, leading=20, spaceAfter=10)
            code_style = ParagraphStyle("Code", parent=styles["Code"], fontName="Courier", fontSize=9, leading=14, spaceAfter=8, leftIndent=20)

            elements = []
            lines = content.split("\n")
            in_code_block = False
            code_lines = []

            for line in lines:
                if line.strip().startswith("```"):
                    if in_code_block:
                        code_text = html.escape("\n".join(code_lines))
                        elements.append(Paragraph(code_text.replace("\n", "<br/>"), code_style))
                        elements.append(Spacer(1, 0.3 * cm))
                        in_code_block = False
                        code_lines = []
                    else:
                        in_code_block = True
                        code_lines = []
                    continue

                if in_code_block:
                    code_lines.append(line)
                    continue

                heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
                if heading_match:
                    level = len(heading_match.group(1))
                    text = html.escape(heading_match.group(2).strip())
                    if level == 1:
                        elements.append(Paragraph(text, title_style))
                    elif level == 2:
                        elements.append(Paragraph(text, h1_style))
                    elif level == 3:
                        elements.append(Paragraph(text, h2_style))
                    else:
                        elements.append(Paragraph(text, h3_style))
                    elements.append(Spacer(1, 0.2 * cm))
                    continue

                if not line.strip():
                    continue

                # 列表
                list_match = re.match(r"^[-*+]\s+(.+)$", line)
                if list_match:
                    text = html.escape(list_match.group(1))
                    elements.append(Paragraph(f"• {text}", body_style))
                    continue

                # 普通段落
                text = html.escape(line)
                text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
                text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
                text = re.sub(r"`(.+?)`", r"<font face='Courier' size=9>\1</font>", text)
                elements.append(Paragraph(text, body_style))

            if in_code_block:
                code_text = html.escape("\n".join(code_lines))
                elements.append(Paragraph(code_text.replace("\n", "<br/>"), code_style))

            doc_pdf.build(elements)
            result_content = None  # PDF 已写入文件

        else:
            self.logger.error("convert: 不支持的目标格式: %s", target_format)
            return {"error": f"不支持的目标格式: {target_format}"}

        # 写入输出文件
        if result_content is None:
            # result_content 为 None 表示已在转换分支内直接写入文件（如 docx、PDF）
            self.logger.info("convert: 格式转换完成, output_path=%s, format=%s", output_path, target_format)
            return {
                "path": output_path,
                "format": target_format,
                "message": f"已转换为 {target_format} 格式",
            }
        elif output_path:
            os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(result_content)
            self.logger.info("convert: 格式转换完成, output_path=%s, format=%s", output_path, target_format)
            return {
                "path": output_path,
                "format": target_format,
                "message": f"已转换为 {target_format} 格式",
            }
        else:
            return {
                "content": result_content,
                "format": target_format,
            }

    def analyze(self, params: dict) -> dict:
        """分析 Markdown 文档"""
        path = params.get("path", "")
        if not path:
            self.logger.error("analyze: 缺少文件路径")
            return {"error": "缺少文件路径"}
        if not os.path.exists(path):
            raise FileNotFoundError(path)

        self.logger.info("analyze: 开始分析 Markdown 文档, path=%s", path)

        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        # 统计
        headings = re.findall(r"^#{1,6}\s+.+$", content, re.MULTILINE)
        code_blocks = re.findall(r"```[\s\S]*?```", content)
        links = re.findall(r"\[([^\]]+)\]\(([^)]+)\)", content)
        images = re.findall(r"!\[([^\]]*)\]\(([^)]+)\)", content)

        self.logger.info("analyze: Markdown 文档分析完成, path=%s, 标题数=%d, 代码块数=%d", path, len(headings), len(code_blocks))
        return {
            "file_size": os.path.getsize(path),
            "char_count": len(content),
            "line_count": content.count("\n") + 1,
            "heading_count": len(headings),
            "code_block_count": len(code_blocks),
            "link_count": len(links),
            "image_count": len(images),
        }

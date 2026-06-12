"""Word 文档生成 Helper 函数
封装 python-docx 常用操作，内置专业配色方案
"""

import os

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 专业配色方案（与原 document_design.rs 保持一致）
THEME = {
    "heading1": (0x1F, 0x4E, 0x79),  # 深蓝色
    "heading2": (0x2E, 0x75, 0xB6),  # 中蓝色
    "heading3": (0x5B, 0x9B, 0xD5),  # 浅蓝色
    "table_header_bg": "D6E4F0",
    "table_alt_row_bg": "EDF2F9",
    "table_border": "B4C6E7",
    "accent": (0x2E, 0x75, 0xB6),
}

EAST_ASIAN_FONT = "微软雅黑"
LATIN_FONT = "Arial"

# 如果 python-docx 可用，将元组转换为 RGBColor 对象
if HAS_DOCX:
    THEME["heading1"] = RGBColor(*THEME["heading1"])
    THEME["heading2"] = RGBColor(*THEME["heading2"])
    THEME["heading3"] = RGBColor(*THEME["heading3"])
    THEME["accent"] = RGBColor(*THEME["accent"])


def create_word_doc(title=None, page_size="a4", author=""):
    """创建一个预配置好专业样式的 Word 文档对象

    Args:
        title: 文档标题（可选）
        page_size: 页面尺寸 "a4" 或 "letter"
        author: 文档作者

    Returns:
        Document: 预配置好的 python-docx Document 对象

    示例:
        doc = create_word_doc(title="季度报告", author="张三")
        doc.add_paragraph("这是正文内容")
        save_word_doc(doc, "季度报告.docx")
    """
    doc = Document()

    # 设置页面尺寸和边距
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    if page_size == "letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # 设置默认字体
    style = doc.styles["Normal"]
    style.font.name = LATIN_FONT
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # 设置标题样式
    for level, (size, color) in enumerate([
        (Pt(22), THEME["heading1"]),
        (Pt(16), THEME["heading2"]),
        (Pt(14), THEME["heading3"]),
    ], start=1):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = LATIN_FONT
        heading_style.font.size = size
        heading_style.font.bold = True
        heading_style.font.color.rgb = color

    # 添加标题
    if title:
        doc.core_properties.title = title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run.font.color.rgb = THEME["heading1"]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if author:
        doc.core_properties.author = author

    return doc


def save_word_doc(doc, filename, working_dir=""):
    """保存 Word 文档到工作目录

    Args:
        doc: python-docx Document 对象
        filename: 文件名（如 "报告.docx"）
        working_dir: 工作目录路径

    Returns:
        str: 保存的文件绝对路径
    """
    output_path = os.path.join(working_dir, filename) if working_dir else filename
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

    return output_path

"""公共样式和工具函数
从原 document_design.rs 配色方案迁移，为所有文档类型提供统一的配色方案
"""

# 专业配色方案（与原 document_design.rs 保持一致）
THEME_COLORS = {
    # 标题颜色
    "heading1": "1F4E79",       # 深蓝色
    "heading2": "2E75B6",       # 中蓝色
    "heading3": "5B9BD5",       # 浅蓝色
    # 表格颜色
    "table_header_bg": "D6E4F0",    # 表头蓝色背景
    "table_alt_row_bg": "EDF2F9",   # 交替行浅蓝色背景
    "table_border": "B4C6E7",       # 蓝灰色边框
    # 强调色
    "accent": "2E75B6",         # 中蓝色强调
    # 文字颜色
    "text_primary": "333333",   # 主文字色
    "text_secondary": "666666", # 次要文字色
    # 背景
    "bg_light": "F5F7FA",       # 浅灰背景
    "bg_white": "FFFFFF",       # 白色背景
}

# 字体配置
FONT_CONFIG = {
    "east_asian": "微软雅黑",
    "latin": "Arial",
    "mono": "Consolas",
}

# PPT 配色方案
PPT_COLOR_SCHEMES = {
    "ocean": {
        "primary": "065A82",
        "secondary": "1C7293",
        "accent": "21295C",
    },
    "midnight": {
        "primary": "1E2761",
        "secondary": "CADCFC",
        "accent": "FFFFFF",
    },
    "forest": {
        "primary": "2C5F2D",
        "secondary": "97BC62",
        "accent": "F5F5F5",
    },
    "coral": {
        "primary": "F96167",
        "secondary": "F9E795",
        "accent": "2F3C7E",
    },
    "charcoal": {
        "primary": "36454F",
        "secondary": "F2F2F2",
        "accent": "212121",
    },
}


def apply_theme(doc, theme_name="default"):
    """对文档应用配色方案

    Args:
        doc: 文档对象（python-docx Document 等）
        theme_name: 配色方案名称

    Returns:
        应用了配色方案的文档对象
    """
    # 当前仅支持 default 主题，后续可扩展
    return doc


def add_styled_table(doc, headers, rows, style_name="Table Grid"):
    """向文档添加专业样式的表格

    Args:
        doc: python-docx Document 对象
        headers: 表头列表
        rows: 数据行列表（二维列表）
        style_name: 表格样式名称

    Returns:
        创建的表格对象
    """
    try:
        from docx.shared import RGBColor, Pt
    except ImportError:
        # python-docx 未安装时，创建无样式的表格
        table = doc.add_table(rows=1 + len(rows), cols=len(headers), style=style_name)
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, cell_data in enumerate(row_data):
                row_cells[col_idx].text = str(cell_data) if cell_data is not None else ""
        return table

    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style=style_name)

    # 添加表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        # 表头样式：蓝色背景 + 粗体
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(THEME_COLORS["heading1"])
                run.font.size = Pt(11)

    # 添加数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data) if cell_data is not None else ""
            for paragraph in row_cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table
